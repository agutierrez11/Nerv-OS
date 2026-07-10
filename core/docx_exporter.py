import os
import re
import unicodedata
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_brackets(text):
    if not text:
        return ""
    # Remove Obsidian wiki-links [[Link]] or [[Link|Text]]
    text = re.sub(r'\[\[([^\]\|]+)\|([^\]]+)\]\]', r'\2', text)
    text = re.sub(r'\[\[([^\]]+)\]\]', r'\1', text)
    return text.strip()

def parse_dossier_markdown(markdown_content):
    """
    Parsea un dossier en markdown y extrae:
    - Nombre del decisor
    - Cargo
    - Correo electrónico
    - Variante D (Asunto, Cuerpo)
    - Variante C (Asunto, Cuerpo)
    - Variante I/S (Asunto, Cuerpo)
    """
    # 1. Extraer datos del tomador de decisiones (Tabla en A QUIÉN HABLARLE)
    nombre = ""
    cargo = ""
    email = ""
    
    nombre_match = re.search(r'\|\s*\*\*Nombre\*\*\s*\|\s*([^\|]+)\|', markdown_content)
    if nombre_match:
        nombre = clean_brackets(nombre_match.group(1))
        
    cargo_match = re.search(r'\|\s*\*\*Cargo\*\*\s*\|\s*([^\|]+)\|', markdown_content)
    if cargo_match:
        cargo = clean_brackets(cargo_match.group(1))
        
    email_match = re.search(r'\|\s*\*\*Correo Electrónico\*\*\s*\|\s*([^\|]+)\|', markdown_content)
    if email_match:
        email_raw = email_match.group(1).strip()
        # Limpiar si no fue detectado o tiene formato bold
        email_raw = email_raw.replace("**", "").strip()
        if "No detectado" in email_raw or "sin créditos" in email_raw.lower():
            email = "No detectado (Buscar en LinkedIn)"
        else:
            email = clean_brackets(email_raw)

    # 2. Extraer Variantes de Correo
    variants = {}
    
    # Patrón general para capturar secciones del tipo "#### Variante [Letra]" hasta la siguiente Variante, separador "---" o final del archivo.
    pattern_var = r'####\s*Variante\s*([A-Za-z/]+).*?\n(.*?)(?=\n####|\n---|#|\Z)'
    matches = re.finditer(pattern_var, markdown_content, re.DOTALL | re.IGNORECASE)
    
    for m in matches:
        var_name = m.group(1).strip() # Ej: "D", "C", "I/S"
        var_text = m.group(2).strip()
        
        # Extraer Asunto
        subject_match = re.search(r'\*\*Asunto:\*\*\s*(.*)', var_text, re.IGNORECASE)
        if not subject_match:
            subject_match = re.search(r'Asunto:\s*(.*)', var_text, re.IGNORECASE)
            
        subject = subject_match.group(1).strip() if subject_match else "Propuesta de valor"
        subject = clean_brackets(subject)
        
        # El cuerpo es el texto remanente quitando la línea del asunto
        body = var_text
        if subject_match:
            body = body.replace(subject_match.group(0), "").strip()
            
        # Limpiar saltos de línea iniciales/finales, backticks de formato y wikilinks
        body = body.strip('` \n')
        body = clean_brackets(body)
        
        variants[var_name] = {
            "subject": subject,
            "body": body
        }
        
    return {
        "nombre": nombre,
        "cargo": cargo,
        "email": email,
        "variants": variants
    }

def convert_dossier_to_docx(markdown_path: Path, output_dir: str = "/home/antonio/Desktop/Correos_Prospeccion"):
    """
    Lee un archivo de dossier en markdown, extrae sus datos y genera un Word (.docx) estructurado.
    """
    if not markdown_path.exists():
        raise FileNotFoundError(f"Dossier no encontrado en {markdown_path}")
        
    content = markdown_path.read_text(encoding="utf-8")
    data = parse_dossier_markdown(content)
    
    empresa = markdown_path.stem.replace("_Reporte", "").replace("_", " ")
    
    # Crear documento Word
    doc = Document()
    
    # Configuración de márgenes y fuentes por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Encabezado Premium
    title = doc.add_paragraph()
    run_title = title.add_run(f"CORREOS DE PROSPECCIÓN: {empresa.upper()}")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue
    title.paragraph_format.space_after = Pt(4)
    
    subtitle = doc.add_paragraph()
    run_sub = subtitle.add_run("Generado automáticamente por NERV OS")
    run_sub.font.italic = True
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(100, 116, 139) # Slate Gray
    subtitle.paragraph_format.space_after = Pt(20)
    
    # Ficha del tomador de decisiones
    p_ficha = doc.add_paragraph()
    p_ficha.add_run("🎯 DATOS DEL DESTINATARIO:").bold = True
    p_ficha.paragraph_format.space_after = Pt(6)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Nombre del Contacto:", "Cargo / Posición:", "Correo Electrónico:"]
    values = [
        data["nombre"] or "No especificado (Completar en el correo)",
        data["cargo"] or "No especificado",
        data["email"] or "No encontrado (Buscar en LinkedIn)"
    ]
    
    for idx in range(3):
        row = table.rows[idx]
        row.cells[0].paragraphs[0].add_run(headers[idx]).bold = True
        row.cells[1].paragraphs[0].add_run(values[idx])
        
    doc.add_paragraph("").paragraph_format.space_after = Pt(20)
    
    # Insertar variantes de correo
    variants = data["variants"]
    if not variants:
        # Fallback si no parseó variantes estructuradas: escribir el texto crudo
        p_err = doc.add_paragraph()
        p_err.add_run("⚠️ No se detectaron secciones de correo estructuradas en el dossier. Por favor, copia el texto del dossier original.").italic = True
    else:
        for var_name, var_data in variants.items():
            # Título de variante
            p_var_title = doc.add_paragraph()
            run_var_title = p_var_title.add_run(f"📬 Variante {var_name}")
            run_var_title.font.bold = True
            run_var_title.font.size = Pt(13)
            run_var_title.font.color.rgb = RGBColor(30, 58, 138)
            p_var_title.paragraph_format.space_before = Pt(15)
            p_var_title.paragraph_format.space_after = Pt(6)
            
            # Asunto
            p_subj = doc.add_paragraph()
            p_subj.paragraph_format.left_indent = Pt(15)
            p_subj.paragraph_format.space_after = Pt(8)
            p_subj.add_run("Asunto: ").bold = True
            run_subj_text = p_subj.add_run(var_data["subject"])
            run_subj_text.font.bold = True
            run_subj_text.font.color.rgb = RGBColor(15, 23, 42)
            
            # Caja de cuerpo del correo
            p_box_start = doc.add_paragraph()
            p_box_start.paragraph_format.left_indent = Pt(15)
            p_box_start.paragraph_format.space_after = Pt(0)
            p_box_start.add_run("------------------------------------ CORREO ------------------------------------").font.color.rgb = RGBColor(148, 163, 184)
            
            # Escribir párrafos
            for line in var_data["body"].split('\n'):
                if line.strip():
                    p_line = doc.add_paragraph()
                    p_line.paragraph_format.left_indent = Pt(15)
                    p_line.paragraph_format.space_after = Pt(6)
                    p_line.paragraph_format.line_spacing = 1.15
                    p_line.add_run(line.strip())
                    
            p_box_end = doc.add_paragraph()
            p_box_end.paragraph_format.left_indent = Pt(15)
            p_box_end.paragraph_format.space_after = Pt(15)
            p_box_end.add_run("--------------------------------------------------------------------------------").font.color.rgb = RGBColor(148, 163, 184)
            
    # Guardar
    out_dir_path = Path(output_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)
    
    nfkd_form = unicodedata.normalize('NFKD', empresa)
    only_ascii = nfkd_form.encode('ASCII', 'ignore').decode('utf-8')
    safe_name = re.sub(r'[^\w\-]', '_', only_ascii).strip('_')
    
    file_path = out_dir_path / f"Correo_{safe_name}.docx"
    doc.save(str(file_path))
    return str(file_path)

if __name__ == "__main__":
    # Test simple
    test_md = """
### 🎯 A QUIÉN HABLARLE

| Elemento | Detalle |
|----------|---------|
| **Nombre** | Eduardo Reyna Navarro |
| **Cargo** | [[CFO]] — Under Armour México |
| **Correo Electrónico** | eduardo.reyna@underarmour.com |

#### Variante D (Perfil Dominante)
**Asunto:** 3 fugas de ingresos en tu checkout
`Eduardo`,
Revisamos tu checkout. Encontramos tres problemas...

#### Variante C (Perfil Analítico)
**Asunto:** Análisis de tasa de aprobación
`Eduardo`,
Realizamos un análisis de la infraestructura...
"""
    parsed = parse_dossier_markdown(test_md)
    print(parsed)
